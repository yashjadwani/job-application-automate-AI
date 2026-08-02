"""DOCX parsing + clean CV regeneration (Option B).

Strategy: parse the uploaded CV into a full ordered structure, let the pipeline
tailor the BULLETS, then REGENERATE a clean, single-column, ATS-friendly DOCX
from that structure. We do not edit the original file, so none of its quirks
(hyperlinks, numbered titles, multi-run bold labels, tab stops) can corrupt the
output.

Parsed shape:
{
  "personal": { "name": str, "contact": str, "lines": [str] },
  "sections": [{
     "id": "sec_0", "title": "SKILLS",
     "bullets": [{ "index": 0, "text": "..." }],       # rewritable — pipeline uses this
     "blocks":  [{ "type": "bullet"|"subhead"|"para", "text": "...", "bi": 0? }],
  }]
}
`blocks` preserves the full ordered content for regeneration; `bullets` is the
subset the tailoring pipeline rewrites (bi links a bullet block to its index).
"""

import io
import re
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

_GLYPHS = "•·▪‣◦●○-–*"

_URL_RE = re.compile(r"https?://[^\s)>\]]+", re.I)
_EMAIL_RE = re.compile(r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b")


def _iter_paragraphs(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    yield from cell.paragraphs


def _num_format(doc, p: Paragraph) -> str | None:
    """The Word numbering format of a paragraph ('bullet', 'decimal', …) or None.
    This is how we tell a real bullet (•) from a numbered title (1. 2. 3.)."""
    pPr = p._p.pPr
    if pPr is None:
        return None
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return None
    num_id_el = numPr.find(qn("w:numId"))
    if num_id_el is None:
        return None
    num_id = num_id_el.get(qn("w:val"))
    ilvl_el = numPr.find(qn("w:ilvl"))
    ilvl = ilvl_el.get(qn("w:val")) if ilvl_el is not None else "0"
    try:
        numbering = doc.part.numbering_part.element
    except Exception:
        return None
    abstract_id = None
    for num in numbering.findall(qn("w:num")):
        if num.get(qn("w:numId")) == num_id:
            a = num.find(qn("w:abstractNumId"))
            abstract_id = a.get(qn("w:val")) if a is not None else None
            break
    if abstract_id is None:
        return None
    for anum in numbering.findall(qn("w:abstractNum")):
        if anum.get(qn("w:abstractNumId")) == abstract_id:
            for lvl in anum.findall(qn("w:lvl")):
                if lvl.get(qn("w:ilvl")) == ilvl:
                    fmt = lvl.find(qn("w:numFmt"))
                    return fmt.get(qn("w:val")) if fmt is not None else None
    return None


def _is_heading(p: Paragraph) -> bool:
    """Section header: a Word Heading/Title style, OR a short ALL-CAPS line
    (hand-formatted CVs). Bullets are never headings."""
    style = (p.style.name or "").lower() if p.style else ""
    if style.startswith("heading") or style == "title":
        return True
    text = p.text.strip()
    if not text or len(text) > 60 or len(text.split()) > 6:
        return False
    if text[0] in _GLYPHS:
        return False
    letters = [c for c in text if c.isalpha()]
    return len(letters) >= 2 and all(c.isupper() for c in letters)


def _is_bullet(p: Paragraph, doc) -> bool:
    """A rewritable bullet: a bullet-glyph list item. NOT a numbered title."""
    text = p.text.strip()
    if not text:
        return False
    fmt = _num_format(doc, p)
    if fmt == "bullet":
        return True
    if fmt is not None:            # decimal/roman/etc → numbered title, not a bullet
        return False
    style = (p.style.name or "").lower() if p.style else ""
    if "list bullet" in style:
        return True
    if text[0] in _GLYPHS and text[:2] != "--":   # manual glyph bullet
        return True
    return False


def _bold(p: Paragraph) -> bool:
    return any(r.bold for r in p.runs)


def _strip_glyph(text: str) -> str:
    text = text.strip()
    if text and text[0] in _GLYPHS:
        text = text[1:].strip()
    return text


def _para_links(para: Paragraph) -> list[dict]:
    """Every link in a paragraph: real <w:hyperlink> targets (URL lives in the
    relationships, not the run text) plus bare URLs/emails typed as plain text.
    Returns [{url, text}], de-duped by url."""
    found: list[dict] = []
    rels = para.part.rels
    for h in para._p.findall(qn("w:hyperlink")):
        rid = h.get(qn("r:id"))
        if not rid or rid not in rels:
            continue
        rel = rels[rid]
        if not rel.is_external:
            continue
        anchor = "".join(t.text or "" for t in h.findall(f".//{qn('w:t')}"))
        found.append({"url": rel.target_ref, "text": anchor.strip()})
    text = para.text
    for m in _URL_RE.finditer(text):
        found.append({"url": m.group(0).rstrip(".,);"), "text": ""})
    for m in _EMAIL_RE.finditer(text):
        found.append({"url": f"mailto:{m.group(0)}", "text": m.group(0)})
    seen: set[str] = set()
    out: list[dict] = []
    for link in found:
        if link["url"] not in seen:
            seen.add(link["url"])
            out.append(link)
    return out


def _classify(url: str) -> str:
    if url.lower().startswith("mailto:"):
        return "email"
    host = (urlparse(url).netloc or "").replace("www.", "").lower()
    if "linkedin.com" in host:
        return "linkedin"
    if "github.com" in host or "gitlab.com" in host:
        return "github"
    return "other"


def _is_portfolio(url: str, name: str) -> bool:
    """A personal site: its host carries a token from the person's name."""
    host = (urlparse(url).netloc or "").replace("www.", "").lower()
    tokens = [t.lower() for t in re.split(r"\W+", name or "") if len(t) > 2]
    return any(t in host for t in tokens)


def parse_docx(data: bytes) -> dict:
    doc = Document(io.BytesIO(data))
    sections: list[dict] = []
    current: dict | None = None
    header: list[str] = []
    links: list[dict] = []

    for para in _iter_paragraphs(doc):
        text = para.text.strip()
        if not text:
            continue
        para_links = _para_links(para)
        if _is_heading(para):
            if current:
                sections.append(current)
            current = {"id": f"sec_{len(sections)}", "title": text,
                       "bullets": [], "blocks": []}
            continue
        if current is None:
            header.append(text)
            for link in para_links:
                links.append({**link, "kind": _classify(link["url"]),
                              "scope": "contact", "section_id": None, "bi": None})
            continue
        if _is_bullet(para, doc):
            clean = _strip_glyph(text)
            bi = len(current["bullets"])
            current["bullets"].append({"index": bi, "text": clean})
            current["blocks"].append({"type": "bullet", "text": clean, "bi": bi,
                                      "links": para_links})
        else:
            style = (para.style.name or "").lower() if para.style else ""
            numbered = _num_format(doc, para) is not None or "number" in style
            kind = "subhead" if (numbered or _bold(para)) else "para"
            bi = None
            current["blocks"].append({"type": kind, "text": text, "links": para_links})
        for link in para_links:
            k = _classify(link["url"])
            links.append({**link, "kind": "project" if k == "other" else k,
                          "scope": "project", "section_id": current["id"], "bi": bi})

    if current:
        sections.append(current)

    name = header[0] if header else ""
    for link in links:
        if link["scope"] == "contact" and link["kind"] == "other" \
                and _is_portfolio(link["url"], name):
            link["kind"] = "portfolio"

    return {
        "personal": {
            "name": name,
            "contact": header[1] if len(header) > 1 else "",
            "lines": header,
        },
        "sections": sections,
        "links": links,
    }


# ---------------------------------------------------------------------------
# Clean CV regeneration (single column, ATS-friendly)
# ---------------------------------------------------------------------------
_INK = RGBColor(0x1D, 0x1D, 0x1F)
_GREY = RGBColor(0x60, 0x60, 0x66)
_ACCENT = RGBColor(0x0A, 0x4A, 0x7A)


def _rule(paragraph):
    """A thin bottom border under a paragraph (section-heading underline)."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "0A4A7A")
    borders.append(bottom)
    pPr.append(borders)


def _add_hyperlink(paragraph, url: str, text: str, size_pt: float | None = None):
    """Append a clickable hyperlink run to a paragraph (python-docx has no
    native API — register an external relationship, then build <w:hyperlink>)."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0A4A7A")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    if size_pt is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text or url
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _render_contact(paragraph, contact: str, links: list[dict]):
    """Emit the header contact line, turning the segments that correspond to a
    captured contact link (LinkedIn, GitHub, Portfolio, email) into clickable
    hyperlinks and leaving the rest (phone, location) as plain grey text."""
    def match(seg: str) -> dict | None:
        for link in links:
            if link.get("text") and link["text"] == seg:
                return link
            if link["kind"] == "email" and link["url"].replace("mailto:", "") in seg:
                return link
        return None

    for i, raw in enumerate(contact.split("|")):
        seg = raw.strip()
        if not seg:
            continue
        if i:
            sep = paragraph.add_run("  |  ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = _GREY
        link = match(seg)
        if link:
            _add_hyperlink(paragraph, link["url"], seg, size_pt=9)
        else:
            r = paragraph.add_run(seg)
            r.font.size = Pt(9)
            r.font.color.rgb = _GREY


def _emit_block_links(paragraph, links: list[dict]):
    for i, link in enumerate(links):
        label = link.get("text") or urlparse(link["url"]).netloc.replace("www.", "")
        paragraph.add_run("  " if i == 0 else " · ")
        _add_hyperlink(paragraph, link["url"], label)


def _prose_without_links(text: str, links: list[dict], strip_anchors: bool) -> str:
    """Text with inline URLs removed (they're re-emitted as clickable links).
    Anchor words are stripped only from original prose, never from a rewrite."""
    text = _URL_RE.sub("", text)
    if strip_anchors:
        for link in links:
            if link.get("text"):
                text = text.replace(link["text"], "")
    return re.sub(r"[\s|•·]+$", "", text).strip()


def render_cv_docx(cv: dict, rewritten_bullets: dict[str, list[str]] | None = None) -> bytes:
    """Build a clean single-column DOCX from the parsed CV, substituting the
    rewritten bullets where provided."""
    rewritten_bullets = rewritten_bullets or {}
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = _INK

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(54)
        section.left_margin = section.right_margin = Pt(54)

    personal = cv.get("personal") or {}

    if personal.get("name"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(personal["name"])
        r.bold = True
        r.font.size = Pt(18)
    if personal.get("contact"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        contact_links = [l for l in (cv.get("links") or []) if l.get("scope") == "contact"]
        _render_contact(p, personal["contact"], contact_links)

    for section in cv.get("sections", []):
        sec_id = section["id"]
        new_bullets = rewritten_bullets.get(sec_id, [])

        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run((section.get("title") or "").upper())
        r.bold = True
        r.font.size = Pt(11.5)
        r.font.color.rgb = _ACCENT
        _rule(h)

        blocks = section.get("blocks")
        # Backward-compat: sections stored before Option B only have `bullets`.
        if not blocks:
            blocks = [{"type": "bullet", "text": b["text"], "bi": b["index"]}
                      for b in section.get("bullets", [])]

        for block in blocks:
            btype = block.get("type")
            text = block.get("text", "")
            blinks = block.get("links") or []
            if btype == "bullet":
                bi = block.get("bi")
                rewritten = bi is not None and bi < len(new_bullets)
                if rewritten:
                    text = new_bullets[bi]
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(2)
                bp.add_run(_prose_without_links(text, blinks, strip_anchors=not rewritten))
                _emit_block_links(bp, blinks)
            elif btype == "subhead":
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(6)
                sp.paragraph_format.space_after = Pt(1)
                sp.add_run(_prose_without_links(text, blinks, strip_anchors=True)).bold = True
                _emit_block_links(sp, blinks)
            else:  # para (e.g. the summary)
                pp = doc.add_paragraph()
                pp.paragraph_format.space_after = Pt(3)
                pp.add_run(_prose_without_links(text, blinks, strip_anchors=True))
                _emit_block_links(pp, blinks)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
