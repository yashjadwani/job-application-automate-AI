"""
Phase 0 de-risk spike: DOCX run-level round-trip.

Goal: prove we can rewrite CV bullet text WITHOUT corrupting any formatting
(fonts, bold, sizes, bullet glyphs, spacing). This is the riskiest assumption
in the CV Tailoring Platform export path.

Usage:
    py spikes/docx_roundtrip.py "C:/path/to/your_cv.docx"

Requirements:
    pip install python-docx

Outputs (written next to your input file):
    <stem>_SPIKE_OUTPUT.docx   -> open in Word: text changed, formatting identical
    <stem>_PARSED.json         -> the sections/bullets structure the parser sees

Nothing is searched, uploaded, or modified in place. Your original is read-only;
all edits go to a NEW copy.
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Block iteration (python-docx skips paragraphs inside tables by default).
# We walk top-level paragraphs AND table cells so nothing is missed.
# ---------------------------------------------------------------------------
def iter_block_paragraphs(parent):
    """Yield (paragraph, location) for body paragraphs and table-cell paragraphs."""
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent), "body"
        elif child.tag == qn("w:tbl"):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p, "table"


def is_heading(paragraph):
    style = (paragraph.style.name or "") if paragraph.style else ""
    return style.lower().startswith("heading") or style.lower() == "title"


def is_bullet(paragraph):
    """A paragraph is a list item if it has numbering (numPr) or a List style."""
    style = (paragraph.style.name or "") if paragraph.style else ""
    if "list" in style.lower():
        return True
    pPr = paragraph._p.pPr
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return True
    return False


def run_format(run):
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "font": run.font.name,
        "size": run.font.size.pt if run.font.size else None,
    }


# ---------------------------------------------------------------------------
# The core technique under test: replace text at run level, preserve formatting.
# All text goes into run[0] (keeping its rPr); other runs are blanked but kept.
# numbering/bullet lives in pPr and is never touched.
# ---------------------------------------------------------------------------
def replace_bullet_text(paragraph, new_text):
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def xml_of(elem):
    from lxml import etree
    return etree.tostring(elem, pretty_print=True).decode() if elem is not None else None


def rpr_snapshot(paragraph):
    """Capture the formatting XML (pPr + each run's rPr) to compare before/after."""
    return {
        "pPr": xml_of(paragraph._p.pPr),
        "runs_rPr": [xml_of(r._element.rPr) for r in paragraph.runs],
    }


def main():
    if len(sys.argv) < 2:
        print("ERROR: pass your CV path, e.g.\n"
              '    py spikes/docx_roundtrip.py "C:/path/to/your_cv.docx"')
        sys.exit(1)

    src = Path(sys.argv[1])
    if not src.exists():
        print(f"ERROR: file not found: {src}")
        sys.exit(1)

    doc = Document(str(src))

    # --- 1. Parse structure -------------------------------------------------
    sections = []
    current = {"id": "preamble", "title": "(preamble)", "bullets": []}
    multi_run_bullets = 0
    table_paragraphs = 0
    all_blocks = list(iter_block_paragraphs(doc))

    for para, location in all_blocks:
        if location == "table":
            table_paragraphs += 1
        text = para.text.strip()
        if not text:
            continue
        if is_heading(para):
            if current["bullets"] or current["id"] != "preamble":
                sections.append(current)
            current = {"id": f"sec_{len(sections)}", "title": text, "bullets": []}
        elif is_bullet(para):
            n_runs = len([r for r in para.runs if r.text])
            if n_runs > 1:
                multi_run_bullets += 1
            current["bullets"].append({
                "index": len(current["bullets"]),
                "text": text,
                "n_runs": n_runs,
                "location": location,
            })
    sections.append(current)

    total_bullets = sum(len(s["bullets"]) for s in sections)

    print("=" * 70)
    print(f"PARSED: {src.name}")
    print("=" * 70)
    print(f"  Total blocks scanned : {len(all_blocks)}")
    print(f"  Sections found       : {len(sections)}")
    print(f"  Bullets found        : {total_bullets}")
    print(f"  Paragraphs in tables : {table_paragraphs}")
    print()
    for s in sections:
        if s["bullets"]:
            print(f"  [{s['id']}] {s['title']}  ({len(s['bullets'])} bullets)")

    # --- 2. Risk flags ------------------------------------------------------
    print()
    print("-" * 70)
    print("RISK FLAGS")
    print("-" * 70)
    if total_bullets == 0:
        print("  !! No bullets detected. This CV may use plain paragraphs or")
        print("     text boxes for experience. The parser needs a different")
        print("     heuristic for this layout.")
    if multi_run_bullets:
        print(f"  !! {multi_run_bullets} bullet(s) span multiple runs.")
        print("     Collapsing to one run would lose mid-bullet formatting")
        print("     (e.g. a bold keyword). Acceptable if bullets are uniformly")
        print("     formatted; otherwise rewrite needs run-aware replacement.")
    if table_paragraphs:
        print(f"  !! {table_paragraphs} paragraph(s) live inside tables.")
        print("     Many CVs lay out content in tables. The export matcher must")
        print("     walk tables too (this script's iterator already does).")
    if not (multi_run_bullets or table_paragraphs or total_bullets == 0):
        print("  OK: simple single-run, non-table bullets. Lowest-risk case.")

    # --- 3. Round-trip test: edit text, prove formatting is untouched -------
    print()
    print("-" * 70)
    print("ROUND-TRIP TEST (text changes, formatting must NOT)")
    print("-" * 70)

    edited = 0
    formatting_violations = 0
    for para, location in all_blocks:
        if not para.text.strip() or not is_bullet(para):
            continue
        before = rpr_snapshot(para)
        before_fmt = [run_format(r) for r in para.runs]

        replace_bullet_text(para, "[SPIKE] " + para.text.strip())

        after = rpr_snapshot(para)
        # pPr (bullet/numbering/indent) must be identical
        if before["pPr"] != after["pPr"]:
            formatting_violations += 1
            print(f"  VIOLATION: pPr changed on a bullet in {location}")
        # run[0] formatting must be identical (only its text changed)
        if before_fmt and para.runs:
            b, a = before_fmt[0], run_format(para.runs[0])
            if (b["bold"], b["italic"], b["font"], b["size"]) != \
               (a["bold"], a["italic"], a["font"], a["size"]):
                formatting_violations += 1
                print(f"  VIOLATION: run[0] formatting changed in {location}")
        edited += 1

    print(f"  Bullets edited        : {edited}")
    print(f"  Formatting violations : {formatting_violations}")
    print(f"  RESULT: {'PASS - run-level edit preserved formatting' if formatting_violations == 0 else 'FAIL - see violations above'}")

    # --- 4. Write outputs ---------------------------------------------------
    out_docx = src.with_name(src.stem + "_SPIKE_OUTPUT.docx")
    out_json = src.with_name(src.stem + "_PARSED.json")
    doc.save(str(out_docx))
    out_json.write_text(json.dumps(sections, indent=2, ensure_ascii=False), encoding="utf-8")

    print()
    print("=" * 70)
    print("OUTPUTS")
    print("=" * 70)
    print(f"  Edited copy : {out_docx}")
    print(f"  Parsed JSON : {out_json}")
    print()
    print("  -> Open the edited copy in Microsoft Word. Every bullet should now")
    print("     start with '[SPIKE] ' but look IDENTICAL in font, weight, size,")
    print("     bullet glyph, and spacing. If it does, the export approach holds.")


if __name__ == "__main__":
    main()
